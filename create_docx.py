#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color', 'space', 'shadow']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_page_number(section):
    """Add page numbers to footer"""
    footer = section.footer
    footer.is_linked_to_previous = False

    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        first_line_indent=1.25, space_before=0, space_after=0):
    """Set standard paragraph formatting"""
    paragraph.alignment = alignment
    paragraph.paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def add_formatted_run(paragraph, text, bold=False, size=14):
    """Add formatted run to paragraph"""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return run

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

# ==================== ТИТУЛЬНЫЙ ЛИСТ ====================
# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(p, 'РОССИЙСКАЯ ФЕДЕРАЦИЯ\n', False, 14)
add_formatted_run(p, 'ИРКУТСКАЯ ОБЛАСТЬ\n', False, 14)
add_formatted_run(p, 'ИРКУТСКОЕ РАЙОННОЕ МУНИЦИПАЛЬНОЕ ОБРАЗОВАНИЕ\n', False, 14)
add_formatted_run(p, 'МУНИЦИПАЛЬНОЕ ОБЩЕОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ\n', False, 14)
add_formatted_run(p, '«СРЕДНЯЯ ОБЩЕОБРАЗОВАТЕЛЬНАЯ ШКОЛА ПОСЕЛКА МОЛОДЕЖНЫЙ»\n\n\n', False, 14)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(p, 'РАЗРАБОТКА WEB-ПРИЛОЖЕНИЯ MAKAROVFLOW\n', True, 16)
add_formatted_run(p, 'ДЛЯ ВЕДЕНИЯ ЛИЧНОГО ДНЕВНИКА И САМОАНАЛИЗА\n\n', True, 16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(p, 'ИТОГОВЫЙ ИНДИВИДУАЛЬНЫЙ ПРОЕКТ\n\n\n\n', True, 14)

# Author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_formatted_run(p, 'Автор:\n', False, 14)
add_formatted_run(p, 'Макаров Роман Евгеньевич, 9Г класс\n', False, 14)
add_formatted_run(p, 'МОУ ИРМО «СОШ п. Молодежный»\n\n', False, 14)
add_formatted_run(p, 'Классный руководитель:\n', False, 14)
add_formatted_run(p, 'Канатопцева Виктория Николаевна\n\n', False, 14)
add_formatted_run(p, 'Подпись: ____________________________\n\n\n\n', False, 14)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(p, 'п. Молодежный\n', False, 14)
add_formatted_run(p, 'Иркутский район\n', False, 14)
add_formatted_run(p, '2025 год', False, 14)

# Page break
doc.add_page_break()

# ==================== СОДЕРЖАНИЕ ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(p, 'СОДЕРЖАНИЕ\n\n', True, 14)

contents = [
    ('ВВЕДЕНИЕ', '3'),
    ('ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ РАЗРАБОТКИ WEB-ПРИЛОЖЕНИЙ', '5'),
    ('    1.1 Анализ существующих решений для ведения дневников', '5'),
    ('    1.2 Обзор современных технологий веб-разработки', '7'),
    ('    1.3 Telegram Mini Apps как платформа для приложений', '9'),
    ('ГЛАВА 2. РАЗРАБОТКА ПРИЛОЖЕНИЯ MAKAROVFLOW', '11'),
    ('    2.1 Проектирование архитектуры приложения', '11'),
    ('    2.2 Реализация основного функционала', '14'),
    ('    2.3 Тестирование и развертывание приложения', '18'),
    ('ЗАКЛЮЧЕНИЕ', '21'),
    ('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', '22'),
    ('ПРИЛОЖЕНИЯ', '24'),
]

for item, page in contents:
    p = doc.add_paragraph()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.JUSTIFY, 0)
    add_formatted_run(p, item + ' ' + '.' * (80 - len(item) - len(page)) + ' ' + page, False, 14)

doc.add_page_break()

# ==================== ВВЕДЕНИЕ ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_formatted_run(p, 'ВВЕДЕНИЕ\n\n', True, 14)

intro_text = """В современном мире цифровых технологий все больше людей стремятся к саморазвитию и самопознанию. Одним из эффективных инструментов для этого является ведение личного дневника, который помогает систематизировать мысли, анализировать эмоции и отслеживать прогресс в достижении целей."""

p = doc.add_paragraph()
set_paragraph_format(p)
add_formatted_run(p, intro_text, False, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_formatted_run(p, '\nАктуальность проекта\n\n', True, 14)

actual_text = """По данным исследований, регулярное ведение дневника способствует улучшению психологического состояния, помогает справляться со стрессом и развивает навыки рефлексии. Однако традиционные бумажные дневники имеют ряд недостатков: их легко потерять, сложно систематизировать записи, невозможно быстро найти нужную информацию.

Существующие цифровые решения для ведения дневников часто требуют установки отдельных приложений, регистрации на различных платформах и не всегда удобны в использовании. Большинство популярных приложений являются платными или содержат навязчивую рекламу.

В связи с этим возникла потребность в создании простого, удобного и доступного инструмента для ведения личного дневника, который был бы интегрирован в уже знакомую пользователям среду — мессенджер Telegram."""

p = doc.add_paragraph()
set_paragraph_format(p)
add_formatted_run(p, actual_text, False, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_formatted_run(p, '\nЦель проекта\n\n', True, 14)

goal_text = """Разработать веб-приложение MakarovFlow для ведения личного дневника с функциями самоанализа, интегрированное в Telegram в виде Mini App, доступное для использования на любых устройствах."""

p = doc.add_paragraph()
set_paragraph_format(p)
add_formatted_run(p, goal_text, False, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_formatted_run(p, '\nЗадачи проекта\n\n', True, 14)

tasks_text = """Для достижения поставленной цели необходимо решить следующие задачи:

1. Изучить существующие решения для ведения электронных дневников и выявить их преимущества и недостатки
2. Изучить современные технологии веб-разработки (React, TypeScript, Vite)
3. Освоить работу с Telegram Mini Apps API
4. Спроектировать архитектуру приложения и базу данных
5. Разработать пользовательский интерфейс приложения
6. Реализовать основной функционал: создание записей, категоризация, поиск
7. Внедрить систему анализа настроения и статистики
8. Протестировать приложение на различных устройствах
9. Развернуть приложение на хостинге и интегрировать с Telegram"""

p = doc.add_paragraph()
set_paragraph_format(p)
add_formatted_run(p, tasks_text, False, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_formatted_run(p, '\nОбъект и предмет исследования\n\n', True, 14)

obj_text = """Объект исследования: процесс разработки веб-приложений для личного использования.

Предмет исследования: методы и технологии создания интерактивных веб-приложений, интегрированных в экосистему Telegram."""

p = doc.add_paragraph()
set_paragraph_format(p)
add_formatted_run(p, obj_text, False, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_formatted_run(p, '\nМетоды исследования\n\n', True, 14)

methods_text = """- Анализ существующих решений и технологий
- Проектирование и моделирование
- Программирование и разработка
- Тестирование и отладка
- Опытная эксплуатация"""

p = doc.add_paragraph()
set_paragraph_format(p)
add_formatted_run(p, methods_text, False, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_formatted_run(p, '\nПрактическая значимость\n\n', True, 14)

pract_text = """Разработанное приложение может использоваться для:
- Ведения личного дневника и журнала мыслей
- Отслеживания эмоционального состояния
- Анализа личного прогресса и достижений
- Развития навыков рефлексии и самоанализа

Приложение является полностью бесплатным, работает без подключения к интернету (после первой загрузки) и не требует регистрации, так как использует аутентификацию Telegram."""

p = doc.add_paragraph()
set_paragraph_format(p)
add_formatted_run(p, pract_text, False, 14)

doc.add_page_break()

print("✅ Создан титульный лист, содержание и введение")
print("📄 Сохранение документа...")

# Save document
doc.save('/Users/romanmakarov/mindflow-app/ИТОГОВЫЙ_ИНДИВИДУАЛЬНЫЙ_ПРОЕКТ.docx')
print("✅ Документ сохранен: ИТОГОВЫЙ_ИНДИВИДУАЛЬНЫЙ_ПРОЕКТ.docx")
print("\n⚠️  Обратите внимание:")
print("   - Добавьте остальные главы вручную в Word")
print("   - Добавьте скриншоты в Приложение А")
print("   - Проверьте форматирование")
